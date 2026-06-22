from __future__ import annotations

import copy
import importlib.util
import json
import py_compile

import pytest

from conftest import ROOT


SKILLS = ROOT / "plugins" / "team-skills" / "skills"
RASPISKA = SKILLS / "raspiska-o-poluchenii-deneg" / "scripts" / "build_raspiska.py"
PACKAGE = SKILLS / "remont-dogovor-i-raspiski" / "scripts" / "build_package.py"


def _load(path, name):
    spec = importlib.util.spec_from_file_location(name, path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def test_glue_scripts_compile(tmp_path) -> None:
    for path in (RASPISKA, PACKAGE):
        assert b"\x00" not in path.read_bytes()
        py_compile.compile(str(path), cfile=str(tmp_path / (path.stem + ".pyc")), doraise=True)


# Скрипты самодостаточны и импортируют python-docx на верхнем уровне (как
# build_ds.py), поэтому модули грузим после importorskip — без пакета тесты
# корректно скипаются.
pytest.importorskip("docx")

RASP = _load(RASPISKA, "build_raspiska")
PKG = _load(PACKAGE, "build_package")


def test_scripts_have_no_render_docx_path_import() -> None:
    # Самодостаточность: не должно быть runtime-импорта соседнего скилла по пути.
    for path in (RASPISKA, PACKAGE):
        src = path.read_text(encoding="utf-8")
        assert "render_docx.py" not in src
        assert "import importlib" not in src
        assert "parents[2]" not in src


# Расписка: чистая сборка структуры -----------------------------------------

RASPISKA_CFG = {
    "city": "г. Алматы",
    "date": "01.06.2026",
    "payer": "Заказчик Синтетический",
    "payee": "Исполнитель Синтетический",
    "amount_digits": "500 000",
    "amount_words": "пятьсот тысяч",
    "currency": "тенге",
    "purpose": "аванс за ремонт квартиры",
}


def _rasp(**overrides):
    cfg = copy.deepcopy(RASPISKA_CFG)
    cfg.update(overrides)
    return cfg


def test_raspiska_structure_has_amount_pair_and_parties() -> None:
    blocks = RASP.build_structure(_rasp())["blocks"]
    text = " ".join(b.get("text", "") for b in blocks if b.get("type") in {"title", "para"})
    assert "РАСПИСКА О ПОЛУЧЕНИИ ДЕНЕГ" in text
    assert "500 000 (пятьсот тысяч) тенге" in text  # цифры и пропись вместе
    assert "Заказчик Синтетический" in text and "Исполнитель Синтетический" in text


def test_raspiska_requires_words_with_digits() -> None:
    with pytest.raises(ValueError):
        RASP.build_structure(_rasp(amount_words=""))


def test_raspiska_keeps_empty_field_for_missing_payer() -> None:
    blocks = RASP.build_structure(_rasp(payer=""))["blocks"]
    sig = [b for b in blocks if b["type"] == "signature_table"][0]
    передал = sig["rows"][0][1]
    assert передал.startswith("____")


def test_raspiska_full_settlement_adds_formula() -> None:
    blocks = RASP.build_structure(_rasp(full_settlement=True))["blocks"]
    text = " ".join(b.get("text", "") for b in blocks)
    assert "получены полностью, претензий не имею" in text


# Пакет: инварианты ----------------------------------------------------------

PACKAGE_CFG = {
    "city": "г. Алматы",
    "contract_date": "01.06.2026",
    "contract_number": "1",
    "customer": "Заказчик Синтетический",
    "contractor": "Исполнитель Синтетический",
    "address": "г. Алматы, ул. Примерная, д. 1, кв. 20",
    "work_list_ref": "перечень работ (приложение № 1)",
    "total_digits": "1 200 000",
    "total_words": "один миллион двести тысяч",
    "currency": "тенге",
    "payments": [
        {"label": "аванс", "amount_digits": "400 000", "amount_words": "четыреста тысяч", "date": "01.06.2026"},
        {"label": "окончательный расчёт", "amount_digits": "800 000", "amount_words": "восемьсот тысяч", "date": "20.06.2026"},
    ],
}


def _pkg(**overrides):
    cfg = copy.deepcopy(PACKAGE_CFG)
    cfg.update(overrides)
    return cfg


def test_invariants_pass_on_consistent_config() -> None:
    PKG.check_invariants(_pkg())  # не должно бросать


def test_invariants_reject_sum_mismatch() -> None:
    bad = _pkg()
    bad["payments"][1]["amount_digits"] = "700 000"  # 400k + 700k ≠ 1.2M
    with pytest.raises(ValueError, match="не сходятся"):
        PKG.check_invariants(bad)


def test_invariants_reject_chronology_violation() -> None:
    bad = _pkg()
    bad["payments"][0]["date"] = "31.05.2026"  # раньше договора 01.06.2026
    with pytest.raises(ValueError, match="Хронология"):
        PKG.check_invariants(bad)


def test_contract_structure_holds_single_address_and_parties() -> None:
    text = " ".join(b.get("text", "") for b in PKG.build_contract_structure(_pkg())["blocks"]
                     if b.get("type") in {"title", "heading", "para"})
    assert "г. Алматы, ул. Примерная, д. 1, кв. 20" in text
    assert "Заказчик Синтетический" in text and "Исполнитель Синтетический" in text
    assert "1 200 000 (один миллион двести тысяч) тенге" in text


# Сквозная сборка .docx ------------------------------------------------------

def test_raspiska_build_writes_valid_docx(tmp_path) -> None:
    from docx import Document

    out = tmp_path / "r.docx"
    RASP.build(_rasp(full_settlement=True), out)
    assert out.exists() and out.stat().st_size > 0
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "500 000 (пятьсот тысяч) тенге" in text
    assert "получены полностью" in text


def test_raspiska_cli_roundtrip_and_format_guard(tmp_path) -> None:
    src = tmp_path / "r.json"
    src.write_text(json.dumps(_rasp()), encoding="utf-8")
    assert RASP.main(["--config", str(src), "--out", str(tmp_path / "r.docx")]) == 0
    assert (tmp_path / "r.docx").exists()
    assert RASP.main(["--config", str(src), "--out", str(tmp_path / "r.pdf")]) == 2


def test_package_builds_contract_plus_receipts(tmp_path) -> None:
    from docx import Document

    outputs = PKG.build_package(_pkg(), tmp_path)
    assert len(outputs) == 3  # договор + 2 расписки
    assert all(p.exists() and p.stat().st_size > 0 for p in outputs)

    contract_text = "\n".join(p.text for p in Document(str(outputs[0])).paragraphs)
    assert "ДОГОВОР ОКАЗАНИЯ УСЛУГ ПО РЕМОНТУ № 1" in contract_text

    # единый адрес во всех документах пакета
    for path in outputs:
        text = "\n".join(p.text for p in Document(str(path)).paragraphs)
        assert "ул. Примерная, д. 1, кв. 20" in text

    # последняя расписка — окончательный расчёт, формула полного получения
    last = "\n".join(p.text for p in Document(str(outputs[-1])).paragraphs)
    assert "получены полностью, претензий не имею" in last


def test_package_does_not_write_files_when_invariant_fails(tmp_path) -> None:
    bad = _pkg()
    bad["payments"][1]["amount_digits"] = "700 000"
    with pytest.raises(ValueError):
        PKG.build_package(bad, tmp_path)
    assert not list(tmp_path.glob("*.docx"))  # инвариант проверяется до записи
