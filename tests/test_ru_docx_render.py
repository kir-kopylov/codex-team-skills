from __future__ import annotations

import importlib.util
import json
import py_compile

import pytest

from conftest import ROOT


SKILL_DIR = ROOT / "plugins" / "team-skills" / "skills" / "ru-docx-render"
SCRIPT = SKILL_DIR / "scripts" / "render_docx.py"


def test_render_docx_has_no_nul_bytes_and_compiles(tmp_path) -> None:
    assert b"\x00" not in SCRIPT.read_bytes()
    py_compile.compile(str(SCRIPT), cfile=str(tmp_path / "render_docx.pyc"), doraise=True)


# --- поведенческие тесты рендера -------------------------------------------
# Запускают сам генератор (как тесты remont-smeta-builder и dopsoglasheniya),
# а не только проверяют компиляцию. Модуль импортирует python-docx на верхнем
# уровне, поэтому без пакета тесты корректно скипаются.

pytest.importorskip("docx")


def _load():
    spec = importlib.util.spec_from_file_location("render_docx", SCRIPT)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


MOD = _load()


def test_new_doc_sets_a4_and_standard_margins() -> None:
    # Поля сравниваем в см с округлением: python-docx хранит w:pgMar в твипах,
    # поэтому 2 см возвращается как 2.00025 см. Округления до 0.1 см достаточно,
    # чтобы отличить наш стандарт от дефолтного шаблона (поля 2.54/3.18 см).
    section = MOD.new_doc().sections[0]
    assert round(section.page_width.cm) == 21
    assert round(section.page_height.cm) == 30
    assert round(section.top_margin.cm, 1) == 2.0
    assert round(section.bottom_margin.cm, 1) == 2.0
    assert round(section.left_margin.cm, 1) == 2.5
    assert round(section.right_margin.cm, 1) == 1.5


def test_helpers_emit_expected_text() -> None:
    doc = MOD.new_doc()
    MOD.title(doc, "РАСПИСКА")
    MOD.city_date(doc, "г. Алматы", "01 июня 2026 г.")
    MOD.heading(doc, "Существо расписки")
    MOD.para(doc, "Получено полностью, претензий нет.")
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "РАСПИСКА" in text
    assert "г. Алматы" in text and "01 июня 2026 г." in text
    assert "Получено полностью, претензий нет." in text


def test_signature_table_keeps_empty_fields() -> None:
    doc = MOD.new_doc()
    table = MOD.signature_table(doc, [["Передал:", "____ / ____"], ["Принял:", "____ / ____"]])
    cells = [cell.text for row in table.rows for cell in row.cells]
    assert "Передал:" in cells
    assert "____ / ____" in cells


def test_signature_table_rejects_empty_rows() -> None:
    with pytest.raises(ValueError):
        MOD.signature_table(MOD.new_doc(), [])


def test_render_structure_builds_valid_docx(tmp_path) -> None:
    from docx import Document

    structure = {
        "blocks": [
            {"type": "title", "text": "РАСПИСКА О ПОЛУЧЕНИИ ДЕНЕГ"},
            {"type": "city_date", "city": "г. Алматы", "date": "01 июня 2026 г."},
            {"type": "para", "text": "Сумма: 500 000 (пятьсот тысяч) тенге."},
            {"type": "signature_table", "rows": [["Передал:", "____"], ["Принял:", "____"]]},
        ]
    }
    out = tmp_path / "raspiska.docx"
    MOD.render_structure(structure, out)
    assert out.exists() and out.stat().st_size > 0
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "РАСПИСКА О ПОЛУЧЕНИИ ДЕНЕГ" in text
    assert "пятьсот тысяч" in text


def test_render_structure_rejects_unknown_and_empty(tmp_path) -> None:
    with pytest.raises(ValueError):
        MOD.render_structure({"blocks": [{"type": "footer", "text": "x"}]}, tmp_path / "a.docx")
    with pytest.raises(ValueError):
        MOD.render_structure({"blocks": []}, tmp_path / "b.docx")


def test_markdown_subset_maps_headings_and_paragraphs() -> None:
    structure = MOD.markdown_to_structure("# Заголовок\n\n## Раздел\n\nОбычный абзац.\n")
    assert [b["type"] for b in structure["blocks"]] == ["title", "heading", "para"]


def test_cli_in_out_roundtrip(tmp_path) -> None:
    from docx import Document

    src = tmp_path / "doc.json"
    src.write_text(json.dumps({"blocks": [{"type": "para", "text": "Тело документа."}]}), encoding="utf-8")
    out = tmp_path / "doc.docx"
    assert MOD.main(["--in", str(src), "--out", str(out)]) == 0
    assert out.exists()
    assert "Тело документа." in "\n".join(p.text for p in Document(str(out)).paragraphs)


def test_cli_rejects_non_docx_output(tmp_path) -> None:
    src = tmp_path / "doc.md"
    src.write_text("# Заголовок\n", encoding="utf-8")
    assert MOD.main(["--in", str(src), "--out", str(tmp_path / "doc.pdf")]) == 2
