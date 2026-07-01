#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Сборщик пакета документов по сделке на ремонт: рамочный договор оказания
услуг + расписки по платежам (самодостаточный, как build_ds.py).

Держит сквозные инварианты ДО сборки:
  * итог = сумме платежей (аванс = сумма частей);
  * хронология: ни одна расписка не датирована раньше договора;
  * единый адрес и имена во всех документах (по построению — из одного config).

Вёрстка выполняется по стандарту оформления скилла `ru-docx-render`, но БЕЗ
runtime-зависимости: рендер встроен в этот файл, расписки формируются по образцу
скилла `raspiska-o-poluchenii-deneg` (тот же текст блоков), без импорта его по
пути. Скрипт не сочиняет недостающие данные: неизвестные личные поля остаются
графой «____».

Использование:
    python build_package.py --config sdelka.json --out-dir ./out
"""
from __future__ import annotations

import argparse
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ModuleNotFoundError as exc:  # pragma: no cover - зависит от окружения автора
    raise SystemExit("Нужен пакет python-docx: установите его в author/runtime окружении.") from exc


# --- стандарт оформления (зеркало ru-docx-render, без импорта) --------------
FONT = "Times New Roman"
SIZE_TITLE = Pt(14)
SIZE_BODY = Pt(12)
EMPTY = "____"
MARGIN_TOP = Cm(2)
MARGIN_BOTTOM = Cm(2)
MARGIN_LEFT = Cm(2.5)
MARGIN_RIGHT = Cm(1.5)
TEXT_WIDTH = Cm(17.0)


def _set_font(run, *, bold: bool = False, size=SIZE_BODY) -> None:
    run.font.size = size
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)


def _new_doc():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = SIZE_BODY
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT
    return doc


def _title(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run(text), bold=True, size=SIZE_TITLE)


def _heading(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_font(p.add_run(text), bold=True)


def _para(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_font(p.add_run(text))


def _city_date(doc, city: str, date: str):
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(TEXT_WIDTH, WD_TAB_ALIGNMENT.RIGHT)
    _set_font(p.add_run(f"{city}\t{date}"))


def _drop_table_borders(table) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        borders.append(el)
    table._tbl.tblPr.append(borders)


def _signature_table(doc, rows: list[list[str]]):
    if not rows:
        raise ValueError("signature_table: нужна хотя бы одна строка")
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _drop_table_borders(table)
    for i, row in enumerate(rows):
        for j in range(cols):
            text = row[j] if j < len(row) else ""
            _set_font(table.cell(i, j).paragraphs[0].add_run(text))


def _req_block(block: dict, key: str) -> Any:
    value = block.get(key)
    if value in (None, ""):
        raise ValueError(f"Блок {block.get('type')!r}: не заполнено поле {key!r}")
    return value


_RENDERERS = {
    "title": lambda doc, b: _title(doc, _req_block(b, "text")),
    "heading": lambda doc, b: _heading(doc, _req_block(b, "text")),
    "para": lambda doc, b: _para(doc, _req_block(b, "text")),
    "city_date": lambda doc, b: _city_date(doc, _req_block(b, "city"), _req_block(b, "date")),
    "signature_table": lambda doc, b: _signature_table(doc, _req_block(b, "rows")),
    "spacer": lambda doc, b: doc.add_paragraph(),
}


def render_structure(structure: dict, out: Path | str) -> Path:
    blocks = structure.get("blocks")
    if not isinstance(blocks, list) or not blocks:
        raise ValueError("Структура должна содержать непустой список blocks")
    doc = _new_doc()
    for block in blocks:
        renderer = _RENDERERS.get(block.get("type"))
        if renderer is None:
            raise ValueError(f"Неизвестный тип блока: {block.get('type')!r}")
        renderer(doc, block)
    out = Path(out)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# --- содержание и инварианты ------------------------------------------------

def require(cfg: dict[str, Any], key: str) -> str:
    value = cfg.get(key)
    if value is None or str(value).strip() == "":
        raise ValueError(f"Не заполнено обязательное поле: {key}")
    return str(value)


def field(value: Any) -> str:
    text = "" if value is None else str(value).strip()
    return text or EMPTY


def amount(digits: Any, words: Any, currency: str) -> str:
    return f"{digits} ({words}) {currency}"


def parse_amount(raw: str, where: str) -> int:
    digits = re.sub(r"\D", "", str(raw))
    if not digits:
        raise ValueError(f"Не удалось разобрать сумму ({where}): {raw!r}")
    return int(digits)


def parse_date(raw: str, where: str) -> datetime:
    try:
        return datetime.strptime(str(raw).strip(), "%d.%m.%Y")
    except ValueError as exc:
        raise ValueError(f"Дата ({where}) ожидается в формате ДД.ММ.ГГГГ: {raw!r}") from exc


def check_invariants(cfg: dict[str, Any]) -> None:
    """Бросает ValueError с понятным текстом при нарушении инварианта —
    программный эквивалент «вынести расхождение пользователю»."""
    require(cfg, "customer")
    require(cfg, "contractor")
    require(cfg, "address")
    payments = cfg.get("payments")
    if not isinstance(payments, list) or not payments:
        raise ValueError("Нужен непустой список payments")

    # 1) итог = сумма частей
    total = parse_amount(require(cfg, "total_digits"), "total_digits")
    parts = sum(parse_amount(p.get("amount_digits"), f"payments[{i}].amount_digits")
                for i, p in enumerate(payments))
    if parts != total:
        raise ValueError(
            f"Суммы не сходятся: сумма платежей {parts} ≠ итог {total}. "
            "Расхождение нужно вынести пользователю, а не подгонять."
        )

    # 2) хронология: платёж/расписка не раньше договора
    contract_dt = parse_date(require(cfg, "contract_date"), "contract_date")
    for i, p in enumerate(payments):
        pay_dt = parse_date(p.get("date"), f"payments[{i}].date")
        if pay_dt < contract_dt:
            raise ValueError(
                f"Хронология нарушена: платёж {i + 1} датирован {p.get('date')} "
                f"раньше договора ({cfg['contract_date']}). Вынести пользователю."
            )


def build_contract_structure(cfg: dict[str, Any]) -> dict:
    city = require(cfg, "city")
    contract_date = require(cfg, "contract_date")
    number = field(cfg.get("contract_number"))
    customer = require(cfg, "customer")
    contractor = require(cfg, "contractor")
    address = require(cfg, "address")
    work_list = field(cfg.get("work_list_ref"))
    currency = require(cfg, "currency")
    total = amount(require(cfg, "total_digits"), require(cfg, "total_words"), currency)

    blocks: list[dict] = [
        {"type": "title", "text": f"ДОГОВОР ОКАЗАНИЯ УСЛУГ ПО РЕМОНТУ № {number}"},
        {"type": "spacer"},
        {"type": "city_date", "city": city, "date": contract_date},
        {"type": "heading", "text": "1. Предмет договора"},
        {"type": "para", "text": (
            f"Исполнитель {contractor} обязуется выполнить работы по ремонту по адресу: "
            f"{address}, в объёме согласно перечню работ ({work_list}), а Заказчик "
            f"{customer} — принять и оплатить их."
        )},
        {"type": "heading", "text": "2. Цена и порядок оплаты"},
        {"type": "para", "text": f"Общая стоимость работ составляет {total}. Оплата производится в следующем порядке:"},
    ]
    for i, p in enumerate(cfg["payments"], start=1):
        line = amount(p.get("amount_digits"), p.get("amount_words"), currency)
        blocks.append({"type": "para", "text": (
            f"{i}) {field(p.get('label'))}: {line} в срок до {field(p.get('date'))}."
        )})
    blocks += [
        {"type": "heading", "text": "3. Сроки и приёмка"},
        {"type": "para", "text": (
            f"Срок выполнения работ: с {EMPTY} по {EMPTY}. Приёмка выполненных работ "
            "оформляется актом со ссылкой на настоящий договор и перечень работ."
        )},
        {"type": "heading", "text": "4. Реквизиты и подписи сторон"},
        {"type": "para", "text": f"Заказчик: {customer}, ИИН {EMPTY}, адрес {EMPTY}."},
        {"type": "para", "text": f"Исполнитель: {contractor}, ИИН {EMPTY}, адрес {EMPTY}."},
        {"type": "spacer"},
        {"type": "signature_table", "rows": [
            ["Заказчик:", f"{customer} / {EMPTY}"],
            ["Исполнитель:", f"{contractor} / {EMPTY}"],
        ]},
    ]
    return {"blocks": blocks}


def build_receipt_structure(cfg: dict[str, Any], payment: dict[str, Any], is_last: bool) -> dict:
    """Расписка по образцу скилла raspiska-o-poluchenii-deneg (тот же текст блоков)."""
    number = field(cfg.get("contract_number"))
    currency = require(cfg, "currency")
    payer = field(cfg.get("customer"))
    payee = field(cfg.get("contractor"))
    line = amount(payment.get("amount_digits"), payment.get("amount_words"), currency)
    purpose = f"{field(payment.get('label'))} по адресу {require(cfg, 'address')}"
    basis = f"договор оказания услуг по ремонту от {require(cfg, 'contract_date')} № {number}"

    blocks: list[dict] = [
        {"type": "title", "text": "РАСПИСКА О ПОЛУЧЕНИИ ДЕНЕГ"},
        {"type": "spacer"},
        {"type": "city_date", "city": require(cfg, "city"), "date": field(payment.get("date"))},
        {"type": "para", "text": (
            f"Я, {payee}, получил(а) от {payer} денежные средства в сумме {line} "
            f"в качестве: {purpose}."
        )},
        {"type": "para", "text": f"Основание: {basis}."},
    ]
    if is_last:
        blocks.append({"type": "para", "text": "Денежные средства получены полностью, претензий не имею."})
    blocks.append({"type": "spacer"})
    blocks.append({"type": "signature_table", "rows": [
        ["Передал:", f"{payer} / {EMPTY}"],
        ["Получил:", f"{payee} / {EMPTY}"],
    ]})
    return {"blocks": blocks}


def build_package(cfg: dict[str, Any], out_dir: Path | str) -> list[Path]:
    """Проверяет инварианты и собирает договор + расписки. Возвращает пути."""
    check_invariants(cfg)
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    outputs: list[Path] = []

    contract_path = out_dir / "dogovor.docx"
    render_structure(build_contract_structure(cfg), contract_path)
    outputs.append(contract_path)

    payments = cfg["payments"]
    for i, payment in enumerate(payments):
        structure = build_receipt_structure(cfg, payment, is_last=(i == len(payments) - 1))
        receipt_path = out_dir / f"raspiska-{i + 1}.docx"
        render_structure(structure, receipt_path)
        outputs.append(receipt_path)

    return outputs


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Сборка пакета: договор на ремонт + расписки.")
    parser.add_argument("--config", required=True, help="JSON-config сделки.")
    parser.add_argument("--out-dir", required=True, help="Каталог для .docx-файлов пакета.")
    args = parser.parse_args(argv)

    cfg = json.loads(Path(args.config).read_text(encoding="utf-8"))
    outputs = build_package(cfg, args.out_dir)
    print("Готово: собран пакет документов:")
    for path in outputs:
        print(f"  - {path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
