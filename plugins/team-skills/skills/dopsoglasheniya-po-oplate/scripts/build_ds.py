#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Генератор DOCX для допсоглашений по порядку оплаты.

Использование:
    python build_ds.py config.json output.docx

Скрипт намеренно не содержит приватных реквизитов. Все рабочие формулировки
передаются через config, извлеченный из договора и подтвержденный пользователем.
"""

from __future__ import annotations

import json
import sys
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ModuleNotFoundError as exc:  # pragma: no cover - зависит от окружения автора
    raise SystemExit("Нужен пакет python-docx: установите его в author/runtime окружении.") from exc


FONT = "Times New Roman"
SZ_TITLE = Pt(13)
SZ_BODY = Pt(13)
SZ_SIGN = Pt(12)
RIGHT_TAB = Cm(17.5)


def require(data: dict[str, Any], key: str) -> Any:
    value = data.get(key)
    if value in (None, ""):
        raise ValueError(f"Не заполнено обязательное поле: {key}")
    return value


def suffix(gender: str, male: str, female: str) -> str:
    return male if gender == "m" else female


def set_run_font(run: Any, size: Any, bold: bool = False) -> None:
    run.font.name = FONT
    run.font.size = size
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:cs"), FONT)
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)


def add_paragraph(
    doc: Any,
    text: str = "",
    *,
    align: Any | None = None,
    size: Any = SZ_BODY,
    bold: bool = False,
    right_tab: bool = False,
    keep_together: bool = False,
    keep_with_next: bool = False,
) -> Any:
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    fmt = paragraph.paragraph_format
    if right_tab:
        fmt.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
    fmt.keep_together = keep_together
    fmt.keep_with_next = keep_with_next
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size, bold)
    return paragraph


def landlord_clause(cfg: dict[str, Any], number: str, address: str, contract_date: str, object_kind: str) -> str:
    landlord = cfg["landlord"]
    gender = landlord.get("gender", "m")
    fio = require(landlord, "fio")
    inn = landlord.get("inn", "")
    landlord_type = require(landlord, "type")
    named = f"именуем{suffix(gender, 'ый', 'ая')}"
    acted = f"действующ{suffix(gender, 'ий', 'ая')}"

    tail = (
        f'{named} в дальнейшем «Арендодатель», с другой стороны, а вместе именуемые '
        f'«Стороны», заключили настоящее Дополнительное соглашение N {number} '
        f'к Договору аренды {object_kind} по адресу: {address} от {contract_date} г. '
        f'(далее - «Договор») о нижеследующем:'
    )

    if landlord_type == "self_employed":
        npd_date = require(landlord, "npd_date")
        npd_number = landlord.get("npd_number", "")
        number_part = f" N {npd_number}" if npd_number else ""
        return (
            f"и {fio}, {acted} как физическое лицо, зарегистрированное в качестве "
            f"налогоплательщика налога на профессиональный доход от {npd_date} г."
            f"{number_part} ИНН {inn}, {tail}"
        )

    if landlord_type == "ip":
        ogrnip = require(landlord, "ogrnip")
        ip_date = require(landlord, "ip_date")
        return (
            f"и Индивидуальный предприниматель {fio}, {acted} на основании свидетельства "
            f"ОГРНИП {ogrnip} от {ip_date} г., {tail}"
        )

    if landlord_type == "individual":
        return f"и {fio}, ИНН {inn}, {tail}"

    raise ValueError(f"Неизвестный тип арендодателя: {landlord_type}")


def common_tail(start_number: int) -> list[str]:
    return [
        f"{start_number}. Настоящее Дополнительное соглашение составлено в двух экземплярах "
        "по одному для каждой из Сторон и является неотъемлемой частью Договора.",
        f"{start_number + 1}. Остальные условия вышеуказанного Договора, не затронутые "
        "настоящим Дополнительным соглашением, остаются неизменными и Стороны подтверждают "
        "по ним свои обязательства.",
    ]


def body_postoplata(params: dict[str, Any]) -> list[str]:
    period_start = require(params, "period_start")
    period_end = require(params, "period_end")
    pay_day = require(params, "pay_day")
    premium_months = params.get("premium_months", "июль и август")
    return_month = params.get("return_month", "сентябрь")
    items = [
        "1. Стороны пришли к соглашению временно изменить порядок внесения арендной платы "
        f"на период с {period_start} по {period_end} включительно. В указанный период оплата "
        f"производится ежемесячно, по факту прошедшего месяца, не позднее {pay_day}-го числа "
        "месяца, следующего за расчетным.",
    ]
    if params.get("rent_amount_digits") or params.get("rent_amount_words"):
        require(params, "rent_amount_digits")
        require(params, "rent_amount_words")
        items.append(
            f"Арендная плата составляет {params['rent_amount_digits']} "
            f"({params['rent_amount_words']}) рублей 00 коп. в месяц."
        )
    items.extend(
        [
            f"2. Арендная плата за {premium_months} вносится в полуторном размере за каждый "
            "из указанных месяцев.",
            f"3. Начиная с расчетов за {return_month}, стороны возвращаются к обычному порядку "
            "оплаты, действовавшему до подписания настоящего Соглашения.",
        ]
    )
    next_number = 4
    if params.get("extra_month") or params.get("extra_amount_digits"):
        require(params, "extra_month")
        require(params, "extra_amount_digits")
        require(params, "extra_amount_words")
        items.append(
            f"{next_number}. Стороны согласовали, что в {params['extra_month']} сумма "
            f"арендной платы составляет {params['extra_amount_digits']} "
            f"({params['extra_amount_words']}) рублей 00 копеек."
        )
        next_number += 1
    items.append(
        f"{next_number}. Изменения, внесенные настоящим Дополнительным соглашением, вступают "
        "в силу с момента подписания, если Стороны не согласовали иной момент применения."
    )
    return items + common_tail(next_number + 1)


def body_vychet(params: dict[str, Any]) -> list[str]:
    month = require(params, "month")
    amount_digits = require(params, "amount_digits")
    amount_words = require(params, "amount_words")
    return [
        f"1. Стороны пришли к соглашению о разовом уменьшении арендной платы за {month} "
        f"на сумму {amount_digits} ({amount_words}) рублей 00 копеек.",
        "2. Указанный вычет является разовым и не влияет на размер арендной платы в "
        "последующие периоды.",
        *common_tail(3),
    ]


def body_izmenenie(params: dict[str, Any]) -> list[str]:
    month = require(params, "month")
    amount_digits = require(params, "amount_digits")
    amount_words = require(params, "amount_words")
    return [
        f"1. Стороны пришли к соглашению, что размер арендной платы за {month} составляет "
        f"{amount_digits} ({amount_words}) рублей 00 копеек.",
        "2. Указанное изменение является разовым и не влияет на размер арендной платы в "
        "последующие периоды.",
        *common_tail(3),
    ]


BODIES = {
    "postoplata": body_postoplata,
    "vychet": body_vychet,
    "izmenenie": body_izmenenie,
}


def add_signatures(doc: Any, cfg: dict[str, Any]) -> None:
    landlord = cfg["landlord"]
    tenant = cfg["tenant"]
    landlord_short = landlord.get("fio_short", landlord.get("fio", ""))
    tenant_signer = tenant.get("signer_short", tenant.get("short_name", ""))
    sign_block = cfg.get("sign_block", "vertical")

    add_paragraph(doc, "Подписи сторон:", align=WD_ALIGN_PARAGRAPH.CENTER, size=SZ_SIGN, bold=True, keep_with_next=True)

    if sign_block == "vertical":
        first = add_paragraph(doc, f"Арендодатель: ________________ {landlord_short}", size=SZ_SIGN, keep_with_next=True)
        first.paragraph_format.space_after = Pt(18)
        add_paragraph(doc, f"Арендатор: ___________________ {tenant_signer}", size=SZ_SIGN)
        return

    if sign_block == "two_col":
        for line in (
            "Арендодатель:\tАрендатор:",
            f"_____________ /{landlord_short}\t_____________ / {tenant_signer}",
        ):
            paragraph = add_paragraph(doc, "", size=SZ_SIGN, keep_together=True)
            paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(9.0), WD_TAB_ALIGNMENT.LEFT)
            run = paragraph.add_run(line)
            set_run_font(run, SZ_SIGN)
        return

    if sign_block == "requisites":
        tenant_short = tenant.get("short_name", "")
        inn = landlord.get("inn", "")
        for line in (
            "АРЕНДОДАТЕЛЬ:\tАРЕНДАТОР:",
            f"{landlord.get('fio', '')}\t{tenant_short}",
            f"ИНН {inn}\t{tenant_signer}",
            f"_______________ / {landlord_short}\t_______________ / {tenant_signer}",
        ):
            paragraph = add_paragraph(doc, "", size=SZ_SIGN, keep_together=True)
            paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(9.0), WD_TAB_ALIGNMENT.LEFT)
            run = paragraph.add_run(line)
            set_run_font(run, SZ_SIGN)
        return

    raise ValueError(f"Неизвестный формат подписного блока: {sign_block}")


def build(cfg: dict[str, Any], out_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"].font.size = SZ_BODY

    number = str(require(cfg, "number"))
    contract_date = require(cfg, "contract_date")
    address = require(cfg, "address")
    object_kind = cfg.get("object_kind", "помещения")

    tenant = require(cfg, "tenant")
    tenant_clause = require(tenant, "full_clause")

    add_paragraph(doc, f"ДОПОЛНИТЕЛЬНОЕ СОГЛАШЕНИЕ N {number}", align=WD_ALIGN_PARAGRAPH.CENTER, size=SZ_TITLE, bold=True)
    add_paragraph(
        doc,
        f"к Договору аренды {object_kind} от {contract_date} г.",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=SZ_TITLE,
        bold=True,
    )
    add_paragraph(doc, f"по адресу: {address}", align=WD_ALIGN_PARAGRAPH.CENTER, size=SZ_TITLE, bold=True)

    date_line = add_paragraph(doc, "", right_tab=True)
    run = date_line.add_run(f"{cfg.get('city', 'г. Екатеринбург')}\t{require(cfg, 'sign_date')}")
    set_run_font(run, SZ_BODY)

    add_paragraph(doc)
    add_paragraph(doc, tenant_clause, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=SZ_BODY)
    add_paragraph(doc, landlord_clause(cfg, number, address, contract_date, object_kind), align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    ds_type = require(cfg, "ds_type")
    if ds_type not in BODIES:
        raise ValueError(f"Неизвестный тип ДС: {ds_type}")
    for item in BODIES[ds_type](cfg.get("params", {})):
        add_paragraph(doc, item, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_signatures(doc, cfg)
    doc.save(out_path)


def main(argv: list[str]) -> int:
    if len(argv) != 3:
        print("Использование: python build_ds.py config.json output.docx")
        return 2
    config_path = Path(argv[1])
    output_path = Path(argv[2])
    with config_path.open(encoding="utf-8") as fh:
        cfg = json.load(fh)
    build(cfg, output_path)
    print(f"OK -> {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
