#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Низкоуровневый рендер документов в .docx по русскому стандарту оформления.

Слой рендера: на вход приходит уже готовая структура документа (JSON-блоки или
markdown-подмножество), на выход — .docx с единым оформлением (Times New Roman,
A4, поля 2/2/2.5/1.5 см, пустые графы «____», таблица подписей без рамок).

Скрипт намеренно НЕ сочиняет содержание, НЕ проверяет смысловые инварианты
(суммы, хронологию, реквизиты) и НЕ заполняет пустые графы — это ответственность
верхнего skill. Здесь только вёрстка.

Использование:
    python render_docx.py --in document.json --out document.docx
    python render_docx.py --in document.md   --out document.docx
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ModuleNotFoundError as exc:  # pragma: no cover - зависит от окружения автора
    raise SystemExit("Нужен пакет python-docx: установите его в author/runtime окружении.") from exc


FONT = "Times New Roman"
SIZE_TITLE = Pt(14)
SIZE_BODY = Pt(12)
EMPTY = "____"  # маркер пустой графы; рендер его не заполняет

# Поля страницы (см): верх / низ / лево / право = 2 / 2 / 2.5 / 1.5
MARGIN_TOP = Cm(2)
MARGIN_BOTTOM = Cm(2)
MARGIN_LEFT = Cm(2.5)
MARGIN_RIGHT = Cm(1.5)

# Ширина текстовой области для правого таба: 21 − 2.5 − 1.5 = 17 см.
TEXT_WIDTH = Cm(17.0)


def _set_font(run, *, bold: bool = False, size=SIZE_BODY) -> None:
    """Жёстко проставляет шрифт на run: ascii + hAnsi + cs, чтобы кириллица
    не уезжала на дефолтный шрифт."""
    run.font.size = size
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)


def new_doc():
    """Пустой документ со стандартом: A4, поля, базовый шрифт Times New Roman."""
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = SIZE_BODY
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT
    return doc


def title(doc, text: str):
    """Заголовок документа: по центру, жирный, увеличенный кегль."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run(text), bold=True, size=SIZE_TITLE)
    return p


def heading(doc, text: str):
    """Заголовок раздела: слева, жирный."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_font(p.add_run(text), bold=True)
    return p


def para(doc, text: str):
    """Абзац тела: по ширине."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_font(p.add_run(text))
    return p


def city_date(doc, city: str, date: str):
    """Строка «город …… дата»: город слева, дата справа по правому табу."""
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(TEXT_WIDTH, WD_TAB_ALIGNMENT.RIGHT)
    _set_font(p.add_run(f"{city}\t{date}"))
    return p


def _drop_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        borders.append(el)
    tbl_pr.append(borders)


def signature_table(doc, rows: list[list[str]]):
    """Таблица подписей без видимых рамок. rows — список строк, каждая строка —
    список ячеек. Пустые графы передаются как «____» и переносятся как есть."""
    if not rows:
        raise ValueError("signature_table: нужна хотя бы одна строка")
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _drop_table_borders(table)
    for i, row in enumerate(rows):
        for j in range(cols):
            text = row[j] if j < len(row) else ""
            cell = table.cell(i, j)
            _set_font(cell.paragraphs[0].add_run(text))
    return table


# --- рендер структуры -------------------------------------------------------

def _req(block: dict, key: str) -> Any:
    value = block.get(key)
    if value in (None, ""):
        raise ValueError(f"Блок {block.get('type')!r}: не заполнено поле {key!r}")
    return value


RENDERERS = {
    "title": lambda doc, b: title(doc, _req(b, "text")),
    "heading": lambda doc, b: heading(doc, _req(b, "text")),
    "para": lambda doc, b: para(doc, _req(b, "text")),
    "city_date": lambda doc, b: city_date(doc, _req(b, "city"), _req(b, "date")),
    "signature_table": lambda doc, b: signature_table(doc, _req(b, "rows")),
    "spacer": lambda doc, b: doc.add_paragraph(),
}


def render_structure(structure: dict, out: Path) -> Path:
    """Собирает .docx из структуры {"blocks": [...]} и сохраняет в out."""
    blocks = structure.get("blocks")
    if not isinstance(blocks, list) or not blocks:
        raise ValueError("Структура должна содержать непустой список blocks")
    doc = new_doc()
    for block in blocks:
        if not isinstance(block, dict):
            raise ValueError(f"Блок должен быть объектом, а не {type(block).__name__}")
        renderer = RENDERERS.get(block.get("type"))
        if renderer is None:
            raise ValueError(f"Неизвестный тип блока: {block.get('type')!r}")
        renderer(doc, block)
    out = Path(out)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# --- markdown-подмножество --------------------------------------------------

def markdown_to_structure(text: str) -> dict:
    """Минимальный markdown → структура: «# » → title, «## » → heading,
    остальные непустые строки → para. Пустые строки — разделители."""
    blocks: list[dict] = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("## "):
            blocks.append({"type": "heading", "text": line[3:].strip()})
        elif line.startswith("# "):
            blocks.append({"type": "title", "text": line[2:].strip()})
        else:
            blocks.append({"type": "para", "text": line})
    return {"blocks": blocks}


def load_structure(path: Path) -> dict:
    raw = path.read_text(encoding="utf-8")
    if path.suffix.lower() == ".json":
        data = json.loads(raw)
        if not isinstance(data, dict):
            raise ValueError("JSON-структура должна быть объектом")
        return data
    return markdown_to_structure(raw)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Рендер структуры или markdown в .docx по стандарту.")
    parser.add_argument("--in", dest="inp", required=True, help="Входной файл: .json (структура) или .md (markdown).")
    parser.add_argument("--out", dest="out", required=True, help="Путь к выходному .docx.")
    args = parser.parse_args(argv)

    out = Path(args.out)
    if out.suffix.lower() != ".docx":
        print("Выходной файл должен иметь расширение .docx", file=sys.stderr)
        return 2

    inp = Path(args.inp)
    if not inp.exists():
        print(f"Не найден входной файл: {inp}", file=sys.stderr)
        return 2

    render_structure(load_structure(inp), out)
    print(f"Готово: собран документ {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
