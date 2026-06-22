#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Сборщик расписки о получении денег (самодостаточный, как build_ds.py).

Главный инвариант — сумма цифрами и сумма прописью обязательны вместе;
отсутствующие личные данные остаются графой «____» (скрипт ничего не выдумывает).

Вёрстка выполняется по стандарту оформления скилла `ru-docx-render`
(Times New Roman, A4, поля 2/2/2.5/1.5 см, таблица подписей без рамок),
но БЕЗ runtime-зависимости от него: рендер встроен в этот файл, чтобы скилл
работал и установленным в одиночку. `ru-docx-render` остаётся референсом
стандарта, а не импортируемым модулем.

Использование:
    python build_raspiska.py --config raspiska.json --out raspiska.docx

config (JSON), минимум:
    {
      "city": "г. Алматы",
      "date": "01.06.2026",
      "payer": "Заказчик (ФИО)",
      "payee": "Исполнитель (ФИО)",
      "amount_digits": "500 000",
      "amount_words": "пятьсот тысяч",
      "currency": "тенге",
      "purpose": "аванс за ремонт квартиры",
      "basis": "по договору оказания услуг от 01.06.2026 № 1",  # опционально
      "full_settlement": false                                    # опционально
    }
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ModuleNotFoundError as exc:  # pragma: no cover - зависит от окружения автора
    raise SystemExit("Нужен пакет python-docx: установите его в author/runtime окружении.") from exc


# --- стандарт оформления (зеркало ru-docx-render, без импорта) --------------
FONT = "Times New Roman"
SIZE_TITLE = Pt(14)
SIZE_BODY = Pt(12)
EMPTY = "____"  # маркер пустой графы; рендер его не заполняет
MARGIN_TOP = Cm(2)
MARGIN_BOTTOM = Cm(2)
MARGIN_LEFT = Cm(2.5)
MARGIN_RIGHT = Cm(1.5)
TEXT_WIDTH = Cm(17.0)  # 21 − 2.5 − 1.5 см: ширина текстовой области для правого таба


def _set_font(run, *, bold: bool = False, size=SIZE_BODY) -> None:
    """Жёстко проставляет шрифт на run (ascii + hAnsi + cs), чтобы кириллица
    не уезжала на дефолтный шрифт."""
    run.font.size = size
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)


def _new_doc():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = SIZE_BODY
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT
    return doc


def _title(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run(text), bold=True, size=SIZE_TITLE)


def _heading(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_font(p.add_run(text), bold=True)


def _para(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_font(p.add_run(text))


def _city_date(doc, city: str, date: str):
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(TEXT_WIDTH, WD_TAB_ALIGNMENT.RIGHT)
    _set_font(p.add_run(f"{city}\t{date}"))


def _drop_table_borders(table) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        borders.append(el)
    table._tbl.tblPr.append(borders)


def _signature_table(doc, rows: list[list[str]]):
    if not rows:
        raise ValueError("signature_table: нужна хотя бы одна строка")
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _drop_table_borders(table)
    for i, row in enumerate(rows):
        for j in range(cols):
            text = row[j] if j < len(row) else ""
            _set_font(table.cell(i, j).paragraphs[0].add_run(text))


def _req_block(block: dict, key: str) -> Any:
    value = block.get(key)
    if value in (None, ""):
        raise ValueError(f"Блок {block.get('type')!r}: не заполнено поле {key!r}")
    return value


_RENDERERS = {
    "title": lambda doc, b: _title(doc, _req_block(b, "text")),
    "heading": lambda doc, b: _heading(doc, _req_block(b, "text")),
    "para": lambda doc, b: _para(doc, _req_block(b, "text")),
    "city_date": lambda doc, b: _city_date(doc, _req_block(b, "city"), _req_block(b, "date")),
    "signature_table": lambda doc, b: _signature_table(doc, _req_block(b, "rows")),
    "spacer": lambda doc, b: doc.add_paragraph(),
}


def render_structure(structure: dict, out: Path | str) -> Path:
    """Собирает .docx из структуры {"blocks": [...]} по стандарту оформления."""
    blocks = structure.get("blocks")
    if not isinstance(blocks, list) or not blocks:
        raise ValueError("Структура должна содержать непустой список blocks")
    doc = _new_doc()
    for block in blocks:
        renderer = _RENDERERS.get(block.get("type"))
        if renderer is None:
            raise ValueError(f"Неизвестный тип блока: {block.get('type')!r}")
        renderer(doc, block)
    out = Path(out)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# --- сборка содержания расписки ---------------------------------------------

def require(cfg: dict[str, Any], key: str) -> str:
    value = cfg.get(key)
    if value is None or str(value).strip() == "":
        raise ValueError(f"Не заполнено обязательное поле: {key}")
    return str(value)


def field(value: Any) -> str:
    """Личные данные: возвращает значение или пустую графу «____»."""
    text = "" if value is None else str(value).strip()
    return text or EMPTY


def format_amount(cfg: dict[str, Any]) -> str:
    """«500 000 (пятьсот тысяч) тенге» — цифры и пропись вместе, обе обязательны."""
    digits = require(cfg, "amount_digits")
    words = require(cfg, "amount_words")
    currency = require(cfg, "currency")
    return f"{digits} ({words}) {currency}"


def build_structure(cfg: dict[str, Any]) -> dict:
    """Чистая функция: config -> структура блоков (без рендера)."""
    city = require(cfg, "city")
    date = require(cfg, "date")
    payer = field(cfg.get("payer"))
    payee = field(cfg.get("payee"))
    amount = format_amount(cfg)
    purpose = require(cfg, "purpose")

    body = (
        f"Я, {payee}, получил(а) от {payer} денежные средства в сумме {amount} "
        f"в качестве: {purpose}."
    )

    blocks: list[dict] = [
        {"type": "title", "text": "РАСПИСКА О ПОЛУЧЕНИИ ДЕНЕГ"},
        {"type": "spacer"},
        {"type": "city_date", "city": city, "date": date},
        {"type": "para", "text": body},
    ]

    basis = str(cfg.get("basis") or "").strip()
    if basis:
        blocks.append({"type": "para", "text": f"Основание: {basis}."})

    if cfg.get("full_settlement"):
        blocks.append({"type": "para", "text": "Денежные средства получены полностью, претензий не имею."})

    blocks.append({"type": "spacer"})
    blocks.append({
        "type": "signature_table",
        "rows": [
            ["Передал:", f"{payer} / {EMPTY}"],
            ["Получил:", f"{payee} / {EMPTY}"],
        ],
    })
    return {"blocks": blocks}


def build(cfg: dict[str, Any], out: Path | str) -> Path:
    """Собирает структуру и рендерит .docx (всё внутри этого скрипта)."""
    return render_structure(build_structure(cfg), out)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Сборка расписки о получении денег в .docx.")
    parser.add_argument("--config", required=True, help="JSON-config расписки.")
    parser.add_argument("--out", required=True, help="Путь к выходному .docx.")
    args = parser.parse_args(argv)

    out = Path(args.out)
    if out.suffix.lower() != ".docx":
        print("Выходной файл должен иметь расширение .docx", file=sys.stderr)
        return 2

    cfg = json.loads(Path(args.config).read_text(encoding="utf-8"))
    build(cfg, out)
    print(f"Готово: собрана расписка {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
